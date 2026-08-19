from docx import Document


def export_word(report_text):

    document = Document()

    document.add_heading(
        "Sustainability Report",
        level=1
    )

    paragraphs = report_text.split("\n")

    for paragraph in paragraphs:

        if paragraph.strip():

            document.add_paragraph(
                paragraph
            )

    output_file = (
        "Generated_Sustainability_Report.docx"
    )

    document.save(
        output_file
    )

    return output_file